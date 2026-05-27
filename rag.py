from pathlib import Path
import hashlib
import json
import logging
import re
import warnings

from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.vectorstores import InMemoryVectorStore

from data_access import infer_category_from_path, infer_marker_from_path
from query_planner import QueryPlan, plan_query


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".tsv",
    ".csv",
    ".fasta",
    ".fa",
    ".html",
    ".htm",
    ".docx",
    ".pdf",
    ".xlsx",
}
IGNORED_FILE_NAMES = {
    "README.md",
    "file_catalog.csv",
    "current_inventory.csv",
}

VECTOR_STORE_PATH = "vectorstore.json"
MANIFEST_PATH = "manifest.json"
INDEX_VERSION = 5


# 给 chunk 文本拼一行可读的元信息头，使 embedding 能感知文件类别和来源。
def _chunk_header(metadata: dict, extra: dict | None = None) -> str:
    category = metadata.get("category")
    marker = metadata.get("marker")
    source = metadata.get("source")

    descriptors = []
    if category and category != "other":
        descriptors.append(category)
    if marker and marker != "general":
        descriptors.append(f"{marker} marker")

    head = "Document"
    if descriptors:
        head += " (" + ", ".join(descriptors) + ")"
    if source:
        head += f": {source}"

    if extra:
        extras = [f"{key}={value}" for key, value in extra.items() if value not in (None, "")]
        if extras:
            head += " | " + " | ".join(extras)

    return head + "\n"

logging.getLogger("pypdf").setLevel(logging.ERROR)
warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")


def should_ignore_file(path: Path) -> bool:
    return path.name in IGNORED_FILE_NAMES or path.name.startswith("~$")


# 按常见编码读取普通文本文件，兼容 UTF-8 和日文 Windows 编码。
def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp932", "shift_jis"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


# 读取 HTML 文件并尽量抽取正文文本。
def read_html_file(path: Path) -> str:
    html = read_text_file(path)
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text("\n")
    except ImportError:
        return re.sub(r"<[^>]+>", " ", html)


# 读取 Word 文档中的段落和表格文本。
def read_docx_file(path: Path) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return "\n".join(parts)


# 读取 PDF 每一页能提取出的文本，并标记页码。
def read_pdf_file(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[page {index}]\n{text}")
    return "\n\n".join(pages)


# 读取 Excel 表格，并为 BLAST 结果表额外生成检索友好的摘要。
def read_xlsx_file(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        parts.append(f"[sheet: {sheet.title}]")
        summary = summarize_blast_rows(rows)
        if summary:
            parts.append(summary)
        for row in rows:
            values = [str(value) for value in row if value is not None]
            if values:
                parts.append("\t".join(values))
    return "\n".join(parts)


# 从 BLAST Target 字段里尽量解析出物种二名法名称。
def parse_species_name(target: object) -> str:
    if not target:
        return "Unknown"

    text = str(target)
    text = re.sub(r"^[A-Z]{1,3}_?\d+(?:\.\d+)?_", "", text)
    text = text.replace(",", " ").replace(";", " ")
    parts = text.split("_")

    for index in range(len(parts) - 1):
        genus = parts[index]
        species = parts[index + 1]
        if genus and species and genus[0].isupper() and species.islower():
            return f"{genus} {species}"

    return text[:100]


# 如果表格是 ASV/BLAST 结果，则按总 reads 和物种生成摘要文本。
def summarize_blast_rows(rows: list[tuple]) -> str:
    if not rows:
        return ""

    header = list(rows[0])
    if "#OTU ID" not in header or "Target" not in header or "Identity" not in header:
        return ""

    target_index = header.index("Target")
    identity_index = header.index("Identity")
    sample_indexes = [
        index for index, value in enumerate(header[:target_index])
        if index > 0 and value
    ]

    records = []
    species_totals = {}
    for row in rows[1:]:
        if not row or not row[0]:
            continue

        total_reads = 0
        for index in sample_indexes:
            value = row[index] if index < len(row) else 0
            try:
                total_reads += int(value or 0)
            except (TypeError, ValueError):
                continue

        if total_reads <= 0:
            continue

        otu_id = row[0]
        target = row[target_index] if target_index < len(row) else ""
        identity = row[identity_index] if identity_index < len(row) else ""
        species = parse_species_name(target)
        records.append((total_reads, otu_id, species, identity, target))
        species_totals[species] = species_totals.get(species, 0) + total_reads

    if not records:
        return ""

    lines = [
        "BLAST summary for retrieval.",
        "Main detected species by total reads:",
    ]

    for species, total_reads in sorted(species_totals.items(), key=lambda item: item[1], reverse=True)[:15]:
        lines.append(f"- {species}: total_reads={total_reads}")

    lines.append("Top ASVs by total reads:")
    for total_reads, otu_id, species, identity, target in sorted(records, reverse=True)[:20]:
        lines.append(
            f"- {otu_id}: species={species}; total_reads={total_reads}; "
            f"identity={identity}; target={target}"
        )

    return "\n".join(lines)


# 根据文件扩展名选择合适的文本抽取方式。
def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".tsv", ".csv", ".fasta", ".fa"}:
        return read_text_file(path)
    if suffix in {".html", ".htm"}:
        return read_html_file(path)
    if suffix == ".docx":
        return read_docx_file(path)
    if suffix == ".pdf":
        return read_pdf_file(path)
    if suffix == ".xlsx":
        return read_xlsx_file(path)
    return ""


# 根据整理后的目录名推断文档类别：论文、表格、报告或序列。
def infer_category(relative_path: str) -> str:
    return infer_category_from_path(Path(relative_path))


# 根据路径推断文档属于 COI、gPlant，还是通用资料。
def infer_marker(relative_path: str) -> str:
    return infer_marker_from_path(Path(relative_path))


# 扫描 data 目录，把支持的文件加载成 LangChain Document。
def load_documents(data_dir: Path) -> list[Document]:
    documents = []
    for path in sorted(data_dir.rglob("*")):
        if should_ignore_file(path):
            continue
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        try:
            text = extract_text(path).strip()
        except Exception as exc:
            warnings.warn(f"Skipping unreadable file {path}: {exc}", RuntimeWarning)
            continue
        if not text:
            continue

        relative_path = str(path.relative_to(data_dir))
        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": relative_path,
                    "category": infer_category(relative_path),
                    "marker": infer_marker(relative_path),
                    "file_type": path.suffix.lower(),
                },
            )
        )
    return documents


# 按表格行切分文档，并保留表头，避免表格结构被普通字符切分破坏。
def split_table_document(
    document: Document,
    rows_per_chunk: int = 40,
    overlap_rows: int = 5,
) -> list[Document]:
    lines = [line for line in document.page_content.splitlines() if line.strip()]
    if not lines:
        return []

    chunks = []
    header = lines[0]
    rows = lines[1:]
    step = rows_per_chunk - overlap_rows
    if step <= 0:
        raise ValueError("rows_per_chunk must be larger than overlap_rows")

    for start in range(0, len(rows), step):
        row_block = rows[start:start + rows_per_chunk]
        if not row_block:
            continue

        metadata = dict(document.metadata)
        metadata["row_start"] = start + 1
        body = "\n".join([header] + row_block)
        chunk_text = _chunk_header(metadata, {"row_start": metadata["row_start"]}) + body
        chunks.append(Document(page_content=chunk_text, metadata=metadata))

    return chunks


# 按 FASTA 的每条序列记录切分，避免把一条序列切断。
def split_fasta_document(document: Document) -> list[Document]:
    chunks = []
    current_header = None
    current_lines = []

    # 将当前正在读取的 FASTA 记录写入 chunk 列表。
    def flush_record():
        if not current_header:
            return
        sequence = "".join(current_lines)
        metadata = dict(document.metadata)
        metadata["sequence_id"] = current_header[1:].split()[0]
        body = f"{current_header}\n{sequence}"
        content = _chunk_header(metadata, {"sequence_id": metadata["sequence_id"]}) + body
        chunks.append(Document(page_content=content, metadata=metadata))

    for line in document.page_content.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith(">"):
            flush_record()
            current_header = line
            current_lines = []
        else:
            current_lines.append(line)

    flush_record()
    return chunks


# 生成数据文件清单和指纹，用于判断本地向量索引是否过期。
def build_manifest(data_dir: Path) -> dict:
    files = []
    for path in sorted(data_dir.rglob("*")):
        if should_ignore_file(path):
            continue
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        stat = path.stat()
        files.append(
            {
                "path": str(path.relative_to(data_dir)),
                "size": stat.st_size,
                "mtime_ns": stat.st_mtime_ns,
            }
        )

    payload = json.dumps(files, ensure_ascii=False, sort_keys=True)
    return {
        "index_version": INDEX_VERSION,
        "fingerprint": hashlib.sha256(payload.encode("utf-8")).hexdigest(),
        "files": files,
    }


# 按不同文件类型选择切分策略，生成最终进入向量库的 chunks。
def split_documents(
    documents: list[Document],
    chunk_size: int = 1400,
    chunk_overlap: int = 200,
) -> list[Document]:
    chunks = []
    step = chunk_size - chunk_overlap
    if step <= 0:
        raise ValueError("chunk_size must be larger than chunk_overlap")

    for document in documents:
        file_type = document.metadata.get("file_type")
        if file_type in {".tsv", ".csv", ".xlsx"}:
            chunks.extend(split_table_document(document))
            continue
        if file_type in {".fasta", ".fa"}:
            chunks.extend(split_fasta_document(document))
            continue

        text = document.page_content
        for start in range(0, len(text), step):
            body = text[start:start + chunk_size].strip()
            if not body:
                continue

            metadata = dict(document.metadata)
            metadata["chunk_start"] = start
            chunk_text = _chunk_header(metadata, {"chunk_start": start}) + body
            chunks.append(Document(page_content=chunk_text, metadata=metadata))

    return chunks


# 用 embedding 模型把 chunks 构建成 LangChain 内存向量库。
def build_vector_store(chunks: list[Document], embeddings) -> InMemoryVectorStore:
    if not chunks:
        raise ValueError("No document chunks found. Add .txt or .md files under data/.")
    return InMemoryVectorStore.from_documents(chunks, embeddings)


# 如果索引未过期则加载本地索引，否则重新构建并保存。
def build_or_load_vector_store(data_dir: Path, index_dir: Path, embeddings, rebuild: bool = False):
    index_dir.mkdir(parents=True, exist_ok=True)
    vector_path = index_dir / VECTOR_STORE_PATH
    manifest_path = index_dir / MANIFEST_PATH
    current_manifest = build_manifest(data_dir)

    if not rebuild and vector_path.exists() and manifest_path.exists():
        saved_manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        if (
            saved_manifest.get("index_version") == current_manifest["index_version"]
            and saved_manifest.get("fingerprint") == current_manifest["fingerprint"]
        ):
            return {
                "vector_store": InMemoryVectorStore.load(str(vector_path), embeddings),
                "rebuilt": False,
                "documents_count": len(current_manifest["files"]),
                "chunks_count": None,
            }

    documents = load_documents(data_dir)
    chunks = split_documents(documents)
    vector_store = build_vector_store(chunks, embeddings)
    vector_store.dump(str(vector_path))
    manifest_path.write_text(
        json.dumps(current_manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    return {
        "vector_store": vector_store,
        "rebuilt": True,
        "documents_count": len(documents),
        "chunks_count": len(chunks),
    }


# 让 LLM 把问题改写成几个不同角度的检索 query，提高召回。无 LLM 时回退原问题。
def expand_queries(question: str, chat_model, n: int = 2) -> list[str]:
    if chat_model is None or n <= 0:
        return [question]

    prompt = (
        f"Rewrite the following question into {n} alternative search queries. "
        "Each rewrite should explore a different angle, synonym, or level of specificity, "
        "while keeping concrete identifiers (sample IDs, ASV IDs, species names, "
        "marker names like COI / gPlant / rbcL) unchanged. "
        "Output one query per line. No numbering, no quotes, no explanation.\n\n"
        f"Question: {question}"
    )
    try:
        response = chat_model.invoke(prompt)
    except Exception:
        return [question]

    text = getattr(response, "content", None) or str(response)
    rewrites = []
    for line in text.splitlines():
        cleaned = line.strip().lstrip("-•·*0123456789.) ").strip()
        if cleaned and cleaned != question and cleaned not in rewrites:
            rewrites.append(cleaned)
        if len(rewrites) >= n:
            break
    return [question] + rewrites


# 在统一的候选池中按 RRF 合并多 query 的结果，并对 plan 偏好做软加权。
def search_with_plan(
    vector_store: InMemoryVectorStore,
    question: str,
    plan: QueryPlan,
    k: int,
    chat_model=None,
):
    queries = expand_queries(question, chat_model)
    pool_size = max(k * 4, 30)

    rrf_scores: dict[tuple, float] = {}
    doc_lookup: dict[tuple, Document] = {}

    for query in queries:
        try:
            candidates = vector_store.similarity_search(query, k=pool_size)
        except Exception:
            continue
        for rank, document in enumerate(candidates):
            key = (
                document.metadata.get("source"),
                document.metadata.get("chunk_start"),
                document.metadata.get("row_start"),
                document.metadata.get("sequence_id"),
                document.page_content[:120],
            )
            rrf_scores[key] = rrf_scores.get(key, 0.0) + 1.0 / (60 + rank)
            doc_lookup.setdefault(key, document)

    if not rrf_scores:
        return []

    allowed_markers = set(plan.markers)
    if "all" in allowed_markers:
        allowed_markers = {"COI", "gPlant", "general"}
    else:
        allowed_markers.add("general")
    plan_categories = set(plan.categories)

    boosted = []
    for key, base_score in rrf_scores.items():
        document = doc_lookup[key]
        marker = document.metadata.get("marker", "")
        category = document.metadata.get("category", "")

        bonus = 0.0
        if marker in allowed_markers:
            bonus += 0.04
        if category in plan_categories:
            bonus += 0.03
        # 论文 / 笔记类内容对几乎所有问题都有背景价值，给一点常驻偏置。
        if category == "knowledge":
            bonus += 0.015

        boosted.append((base_score + bonus, document))

    boosted.sort(key=lambda item: -item[0])
    return [document for _, document in boosted[:k]]


# 把检索到的 chunks 组装成提供给 LLM 的上下文文本。
def format_context(documents: list[Document]) -> str:
    parts = []
    for index, document in enumerate(documents, start=1):
        source = document.metadata.get("source", "unknown")
        parts.append(f"[{index}] Source: {source}\n{document.page_content}")
    return "\n\n".join(parts)


# 构建 RAG 回答链：Prompt -> Chat model -> 字符串输出。
def build_rag_chain(chat_model):
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are a careful, document-grounded research assistant. "
                "Use the provided context as your primary source of truth. "
                "You may synthesize and connect information across multiple context entries, "
                "and you may rephrase, summarize, or reorganize what is in the context. "
                "However, do not introduce facts, numbers, names, species, methods, or conclusions "
                "that are not present in the context. "
                "If the context only partially answers the question, give the partial answer "
                "and clearly state which parts are unclear or missing — "
                "do not refuse outright when useful partial information exists. "
                "Only say the documents do not confirm something when the context truly contains no relevant information. "
                "Answer in the same language as the user's question. "
                "Cite supporting context entries by bracket number, e.g. [1], [2].",
            ),
            (
                "human",
                "Context:\n{context}\n\n"
                "Question: {question}\n\n"
                "Synthesize a clear, helpful answer from the context. "
                "Prefer giving a partial answer with explicit gaps over refusing entirely.",
            ),
        ]
    )
    return prompt | chat_model | StrOutputParser()


# 完整执行一次 RAG：规划问题、加载索引、检索证据、调用模型回答。
def answer_question(
    question: str,
    chat_model,
    embeddings,
    data_dir: Path,
    index_dir: Path,
    k: int = 4,
    rebuild_index: bool = False,
):
    plan = plan_query(question, chat_model=chat_model)
    index = build_or_load_vector_store(data_dir, index_dir, embeddings, rebuild=rebuild_index)
    vector_store = index["vector_store"]
    retrieved_docs = search_with_plan(vector_store, question, plan, k=k, chat_model=chat_model)

    chain = build_rag_chain(chat_model)
    answer = chain.invoke(
        {
            "context": format_context(retrieved_docs),
            "question": question,
        }
    )

    sources = []
    for document in retrieved_docs:
        source = document.metadata.get("source", "unknown")
        if source not in sources:
            sources.append(source)

    return {
        "answer": answer,
        "sources": sources,
        "retrieved_documents": [
            {
                "source": document.metadata.get("source", "unknown"),
                "marker": document.metadata.get("marker"),
                "category": document.metadata.get("category"),
                "chunk_start": document.metadata.get("chunk_start"),
                "content": document.page_content,
            }
            for document in retrieved_docs
        ],
        "retrieved_count": len(retrieved_docs),
        "query_plan": {
            "markers": list(plan.markers),
            "categories": list(plan.categories),
            "mode": plan.mode,
        },
        "index_rebuilt": index["rebuilt"],
        "documents_count": index["documents_count"],
        "chunks_count": index["chunks_count"],
    }


# ---------------------------------------------------------------------------
# Gather + compose pipeline (replaces the if/elif routing in main.py)
# ---------------------------------------------------------------------------


# 从向量库取回 RAG 候选片段，但不直接生成答案。
def gather_rag_chunks(
    question: str,
    chat_model,
    embeddings,
    data_dir: Path,
    index_dir: Path,
    k: int = 8,
    rebuild_index: bool = False,
):
    plan = plan_query(question, chat_model=chat_model)
    index = build_or_load_vector_store(data_dir, index_dir, embeddings, rebuild=rebuild_index)
    vector_store = index["vector_store"]
    retrieved_docs = search_with_plan(vector_store, question, plan, k=k, chat_model=chat_model)
    return {
        "plan": plan,
        "retrieved_docs": retrieved_docs,
        "index_info": index,
    }


# 把 RAG 返回的 chunks 转换成 evidence 条目。
def format_rag_entries(rag_evidence) -> list[dict]:
    if not rag_evidence:
        return []
    entries = []
    for document in rag_evidence["retrieved_docs"]:
        entries.append(
            {
                "kind": "rag",
                "source": document.metadata.get("source", "unknown"),
                "text": document.page_content,
            }
        )
    return entries


# 把所有来源的 evidence 拼成统一编号的文本块，给 composer 用。
def render_evidence_block(entries: list[dict]) -> str:
    if not entries:
        return "(no evidence collected)"
    parts = []
    for index, entry in enumerate(entries, start=1):
        header = f"[{index}] kind={entry['kind']} | source={entry['source']}"
        parts.append(f"{header}\n{entry['text']}")
    return "\n\n".join(parts)


# Composer：拿到所有 gather 出来的证据（结构化 + RAG），让 LLM 写一份最终答案。
def build_compose_chain(chat_model):
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are a careful, document-grounded research assistant. "
                "The user's question has been routed through several backends, and the system has "
                "gathered evidence below from structured analysis tables, FASTA / BLAST data, "
                "sample metadata from thesis / field documents, and retrieved document chunks. "
                "Each evidence entry is numbered [N] with kind and source. "
                "\n\n"
                "Write a single coherent answer in the same language as the user's question. "
                "Use ONLY the evidence below. Do not invent facts not present in it. "
                "When a structured-table evidence entry says a sample / ASV / species is not in a file, "
                "state that fact directly using the column lists or available items shown — "
                "do not refuse to answer. "
                "When evidence is partial, give the partial answer and state the gaps explicitly. "
                "Cite supporting evidence by bracket number, e.g. [1], [2], where applicable. "
                "If the evidence is genuinely unrelated to the question, say so plainly.",
            ),
            (
                "human",
                "Question: {question}\n\nEvidence:\n{evidence}",
            ),
        ]
    )
    return prompt | chat_model | StrOutputParser()


def compose_answer(question: str, entries: list[dict], chat_model) -> str | None:
    if chat_model is None:
        return None
    chain = build_compose_chain(chat_model)
    return chain.invoke(
        {
            "question": question,
            "evidence": render_evidence_block(entries),
        }
    )
